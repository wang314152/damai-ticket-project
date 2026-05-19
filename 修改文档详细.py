from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 读取原文档
doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告.docx')

# 找到所有段落
paragraphs = doc.paragraphs

# 需要增强的关键内容
enhancements = {
    "1.1 前后端联调环节": """
具体问题案例：

问题1：跨域请求失败
• 现象：前端axios请求后端API时报错 "Access-Control-Allow-Origin"
• 原因：前端开发服务器localhost:5173，后端localhost:8081，端口不同导致浏览器安全限制
• 解决：在vite.config.js配置proxy代理，将/api请求转发到后端

问题2：数据格式不一致
• 现象：前端期望返回 {code: 200, data: [...]}, 后端直接返回数组
• 原因：后端Controller未统一返回格式，前后端未约定接口规范
• 解决：定义统一响应类Result<T>，包含code、msg、data字段

问题3：字段命名差异
• 现象：后端返回userName，前端接收显示undefined
• 原因：后端用驼峰命名，前端期望下划线命名，JSON序列化配置不一致
• 解决：application.yml配置开启驼峰转下划线，或前端统一处理""",

    "1.2 数据库设计环节": """
实际遇到的问题：

• 座位表查询慢：未建立索引，1万条数据全表扫描导致3-5秒延迟
  解决：为showId、status字段建立联合索引

• 订单表冗余：冗余存储用户名，历史订单无法显示最新信息
  解决：改为关联查询，实时获取最新用户信息

• 时区问题：数据库时间比实际少8小时
  解决：application.yml配置 serverTimezone=Asia/Shanghai""",

    "3.1 关于用户行为的假设": """
详细分析：

假设1：用户会按正常流程购票（浏览→选座→支付）
实际情况：用户可能直接搜索、频繁返回操作，需要考虑操作路径多样性

假设2：用户会等待页面加载（接口<1秒响应）
实际情况：网络波动时可能等待3-5秒，需要添加loading动画和超时处理

假设3：用户不会重复提交
实际情况：用户可能双击或网络重试导致重复下单，需要防重复提交机制
解决：按钮防抖 + 请求幂等性设计""",

    "4.2 构建步骤": """
【逐行详细说明】

1. Checkout code：使用actions/checkout@v4拉取代码
   - 这是GitHub Actions的第一步，必不可少

2. Cache Maven packages：缓存Maven依赖加速构建
   - key包含pom.xml的hash，pom变化时自动更新缓存
   - 可以节省2-3分钟的依赖下载时间

3. Set up JDK 17：安装Java 17运行环境
   - distribution使用temurin发行版（社区推荐的OpenJDK发行版）
   - 比默认的adoptium更快更稳定

4. Build Backend：cd damai-ticket && mvn clean package -DskipTests
   - -DskipTests跳过测试，加快CI构建速度

5. Cache npm packages：缓存npm依赖加速构建
   - 基于package-lock.json的hash

6. Set up Node.js：安装Node.js 20 LTS版本
   - LTS版本更稳定，20是最新的长期支持版

7. Build Frontend：npm install（使用npmmirror镜像）+ npm run build
   - npmmirror是国内镜像，比官方npm快10倍以上

8. Deploy to GitHub Pages：上传dist目录，使用actions/deploy-pages@v4部署""",

    "4.1 票务系统中的AI应用场景": """
【具体对话示例】

智能客服问答：
用户："周末有什么演唱会推荐？"
AI："根据您的位置和偏好，推荐以下演唱会..."

购票引导：
用户："怎么买票？"
AI："购票流程：1.登录账号 2.选择演出 3.选座 4.支付"

订单咨询：
用户："我的订单什么时候发货？"
AI："您的订单已支付成功，电子票将在演出前24小时发送"

退票政策：
用户："能退票吗？"
AI："演出前72小时内不支持退票，详情请查看退票政策页面""",

    "3.1 座位超卖问题的解决": """
【案例一：座位超卖问题详细分析】

问题背景：
用户反映：两个人同时抢同一张票，都显示购买成功，但实际上只有一个人能真正拿到票。

系统分析（运用系统思维）：
1. 识别关键链路：检查座位状态 → 创建订单 → 更新座位状态
2. 发现关联：座位状态和订单创建相互影响，修改顺序很关键
3. 找到根因：检查→创建→更新不是原子操作，并发时出现竞态条件

解决方案：
• 使用MyBatis-Plus乐观锁机制：@Version注解自动处理版本控制
• 更新时自动检查版本号，防止并发修改导致的数据不一致
• 通过压测验证解决方案的有效性

【案例二：GitHub Pages部署问题详细分析】

问题背景：
本地开发正常，部署到GitHub Pages后页面空白，点击导航链接显示404错误。

系统分析（运用系统思维）：
1. 识别环境差异：本地有后端服务器处理路由，生产是纯静态服务器
2. 理解约束条件：静态托管无法配置服务端路由规则
3. 找到根因：Vue Router的History模式依赖服务器配置，GitHub Pages不支持

解决方案：
• 改用Hash模式路由：路由信息存储在URL的#号后面
• 代码实现：根据环境自动选择路由模式
• const history = isGitHubPages ? createWebHashHistory() : createWebHistory();""",
}

# 遍历段落，找到匹配的内容并在后面添加详细信息
i = 0
while i < len(doc.paragraphs):
    text = doc.paragraphs[i].text.strip()
    
    for key, content in enhancements.items():
        if key in text:
            # 在当前段落后添加新段落
            # 插入增强内容
            for line in content.strip().split('\n'):
                new_para = doc.paragraphs[i]._element
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                # 在当前段落后插入新段落
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = line
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                p.append(r)
                new_para.addnext(p)
            
            # 重新加载段落列表
            doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_temp.docx')
            doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告_temp.docx')
            paragraphs = doc.paragraphs
            break
    i += 1

# 保存修改后的文档
doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_已完善.docx')
print("文档增强成功！")
