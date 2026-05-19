from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 读取原文档
doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告.docx')

def add_heading(doc, text, level=2):
    """添加标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 2 else 12)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return para

def add_text(doc, text, bold=False):
    """添加正文"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para

def add_bullet(doc, text):
    """添加列表项"""
    para = doc.add_paragraph()
    run = para.add_run("• " + text)
    run.font.size = Pt(11)
    return para

# ==================== Day 1 详细补充 ====================
doc.add_paragraph()
add_heading(doc, "【Day 1 详细补充】问题分析深入", level=1)

add_heading(doc, "1.1 前后端联调环节 - 详细案例分析", level=2)

add_text(doc, "【问题1：跨域请求失败】", bold=True)
add_bullet(doc, "现象：前端axios请求后端API时报错 \"Access-Control-Allow-Origin\"")
add_bullet(doc, "原因：前端开发服务器localhost:5173，后端localhost:8081，端口不同导致浏览器安全限制")
add_bullet(doc, "解决：在vite.config.js配置proxy代理，将/api请求转发到后端")
add_text(doc, "关键代码：", bold=True)
add_text(doc, "proxy: { '/api': { target: 'http://localhost:8081', changeOrigin: true } }")

add_text(doc, "")
add_text(doc, "【问题2：数据格式不一致】", bold=True)
add_bullet(doc, "现象：前端期望返回 {code: 200, data: [...]}, 后端直接返回数组")
add_bullet(doc, "原因：后端Controller未统一返回格式，前后端未约定接口规范")
add_bullet(doc, "解决：定义统一响应类Result<T>，包含code、msg、data字段")

add_text(doc, "")
add_text(doc, "【问题3：字段命名差异】", bold=True)
add_bullet(doc, "现象：后端返回userName，前端接收显示undefined")
add_bullet(doc, "原因：后端用驼峰命名，前端期望下划线命名，JSON序列化配置不一致")
add_bullet(doc, "解决：application.yml配置开启驼峰转下划线，或前端统一处理")

add_heading(doc, "1.2 数据库设计环节 - 性能问题", level=2)
add_bullet(doc, "座位表查询慢：未建立索引，1万条数据全表扫描导致3-5秒延迟")
add_bullet(doc, "  解决：为showId、status字段建立联合索引")
add_bullet(doc, "订单表冗余：冗余存储用户名，历史订单无法显示最新信息")
add_bullet(doc, "  解决：改为关联查询，实时获取最新用户信息")
add_bullet(doc, "时区问题：数据库时间比实际少8小时")
add_bullet(doc, "  解决：application.yml配置 serverTimezone=Asia/Shanghai")

add_heading(doc, "3.1 关于用户行为的假设 - 详细分析", level=2)
add_text(doc, "假设1：用户会按正常流程购票（浏览→选座→支付）", bold=True)
add_text(doc, "实际情况：用户可能直接搜索、频繁返回操作，需要考虑操作路径多样性", False)
add_text(doc, "")
add_text(doc, "假设2：用户会等待页面加载（接口<1秒响应）", bold=True)
add_text(doc, "实际情况：网络波动时可能等待3-5秒，需要添加loading动画和超时处理", False)
add_text(doc, "")
add_text(doc, "假设3：用户不会重复提交", bold=True)
add_text(doc, "实际情况：用户可能双击或网络重试导致重复下单，需要防重复提交机制", False)
add_text(doc, "解决：按钮防抖 + 请求幂等性设计", False)

# ==================== Day 3 详细补充 ====================
doc.add_paragraph()
add_heading(doc, "【Day 3 详细补充】CI/CD配置解析", level=1)

add_heading(doc, "4.2 构建步骤 - 逐行详解", level=2)
add_text(doc, "1. Checkout code：使用actions/checkout@v4拉取代码", bold=True)
add_text(doc, "   这是GitHub Actions的第一步，必不可少")
add_text(doc, "")
add_text(doc, "2. Cache Maven packages：缓存Maven依赖加速构建", bold=True)
add_text(doc, "   key包含pom.xml的hash，pom变化时自动更新缓存")
add_text(doc, "   可以节省2-3分钟的依赖下载时间")
add_text(doc, "")
add_text(doc, "3. Set up JDK 17：安装Java 17运行环境", bold=True)
add_text(doc, "   distribution使用temurin发行版（社区推荐的OpenJDK发行版）")
add_text(doc, "   比默认的adoptium更快更稳定")
add_text(doc, "")
add_text(doc, "4. Build Backend：cd damai-ticket && mvn clean package -DskipTests", bold=True)
add_text(doc, "   -DskipTests跳过测试，加快CI构建速度")
add_text(doc, "")
add_text(doc, "5. Cache npm packages：缓存npm依赖加速构建", bold=True)
add_text(doc, "   基于package-lock.json的hash")
add_text(doc, "")
add_text(doc, "6. Set up Node.js：安装Node.js 20 LTS版本", bold=True)
add_text(doc, "   LTS版本更稳定，20是最新的长期支持版")
add_text(doc, "")
add_text(doc, "7. Build Frontend：npm install（使用npmmirror镜像）+ npm run build", bold=True)
add_text(doc, "   npmmirror是国内镜像，比官方npm快10倍以上")
add_text(doc, "")
add_text(doc, "8. Deploy to GitHub Pages：上传dist目录，使用actions/deploy-pages@v4部署", bold=True)

# ==================== Day 4 详细补充 ====================
doc.add_paragraph()
add_heading(doc, "【Day 4 详细补充】AI应用场景", level=1)

add_heading(doc, "4.1 票务系统中的AI应用场景 - 具体对话示例", level=2)
add_text(doc, "【对话1：智能客服问答】", bold=True)
add_text(doc, "用户：\"周末有什么演唱会推荐？\"")
add_text(doc, "AI：\"根据您的位置和偏好，推荐以下演唱会...\"")
add_text(doc, "")
add_text(doc, "【对话2：购票引导】", bold=True)
add_text(doc, "用户：\"怎么买票？\"")
add_text(doc, "AI：\"购票流程：1.登录账号 2.选择演出 3.选座 4.支付\"")
add_text(doc, "")
add_text(doc, "【对话3：订单咨询】", bold=True)
add_text(doc, "用户：\"我的订单什么时候发货？\"")
add_text(doc, "AI：\"您的订单已支付成功，电子票将在演出前24小时发送\"")
add_text(doc, "")
add_text(doc, "【对话4：退票政策】", bold=True)
add_text(doc, "用户：\"能退票吗？\"")
add_text(doc, "AI：\"演出前72小时内不支持退票，详情请查看退票政策页面\"")

# ==================== Day 8 详细补充 ====================
doc.add_paragraph()
add_heading(doc, "【Day 8 详细补充】系统思维案例", level=1)

add_heading(doc, "3.1 座位超卖问题的解决 - 系统思维应用", level=2)
add_text(doc, "【问题背景】", bold=True)
add_text(doc, "用户反映：两个人同时抢同一张票，都显示购买成功，但实际上只有一个人能真正拿到票。")
add_text(doc, "")
add_text(doc, "【系统分析】", bold=True)
add_bullet(doc, "识别关键链路：检查座位状态 → 创建订单 → 更新座位状态")
add_bullet(doc, "发现关联：座位状态和订单创建相互影响，修改顺序很关键")
add_bullet(doc, "找到根因：检查→创建→更新不是原子操作，并发时出现竞态条件")
add_text(doc, "")
add_text(doc, "【解决方案】", bold=True)
add_bullet(doc, "使用MyBatis-Plus乐观锁机制：@Version注解自动处理版本控制")
add_bullet(doc, "更新时自动检查版本号，防止并发修改导致的数据不一致")
add_bullet(doc, "通过压测验证解决方案的有效性")

add_text(doc, "")
add_heading(doc, "3.2 GitHub Pages部署问题 - 系统思维应用", level=2)
add_text(doc, "【问题背景】", bold=True)
add_text(doc, "本地开发正常，部署到GitHub Pages后页面空白，点击导航链接显示404错误。")
add_text(doc, "")
add_text(doc, "【系统分析】", bold=True)
add_bullet(doc, "识别环境差异：本地有后端服务器处理路由，生产是纯静态服务器")
add_bullet(doc, "理解约束条件：静态托管无法配置服务端路由规则")
add_bullet(doc, "找到根因：Vue Router的History模式依赖服务器配置，GitHub Pages不支持")
add_text(doc, "")
add_text(doc, "【解决方案】", bold=True)
add_bullet(doc, "改用Hash模式路由：路由信息存储在URL的#号后面")
add_bullet(doc, "代码实现：根据环境自动选择路由模式")
add_bullet(doc, "const history = isGitHubPages ? createWebHashHistory() : createWebHistory();")

# ==================== 项目总结 ====================
doc.add_paragraph()
add_heading(doc, "【项目总结与展望】", level=1)

add_heading(doc, "一、项目概述", level=2)
add_text(doc, "本项目是一个前后端分离的在线票务系统，模拟大麦网的演出票务购买流程。采用Vue 3 + Spring Boot技术栈，实现了演出浏览、座位选择、订单管理、AI智能客服等核心功能，并搭建了完整的CI/CD自动化部署流程。")

add_heading(doc, "二、技术收获", level=2)
add_text(doc, "2.1 前端技术", bold=True)
add_bullet(doc, "Vue 3 Composition API：掌握组件逻辑复用，告别Options API的繁琐")
add_bullet(doc, "Vue Router：深入理解SPA路由原理（History vs Hash模式）")
add_bullet(doc, "Element Plus：组件库使用和主题定制，提升开发效率")
add_bullet(doc, "Vite：体验到现代构建工具的快速开发体验（热更新秒级响应）")
add_text(doc, "")
add_text(doc, "2.2 后端技术", bold=True)
add_bullet(doc, "Spring Boot：从零搭建企业级应用，理解自动配置原理")
add_bullet(doc, "MyBatis-Plus：掌握ORM框架和乐观锁机制，解决并发问题")
add_bullet(doc, "RESTful API：接口设计规范，统一响应格式code/msg/data")
add_text(doc, "")
add_text(doc, "2.3 DevOps技术", bold=True)
add_bullet(doc, "GitHub Actions：CI/CD流水线配置，实现自动化构建部署")
add_bullet(doc, "GitHub Pages：静态网站托管，免费的演示环境")
add_bullet(doc, "环境配置管理：开发/生产环境分离，配置即代码")

add_heading(doc, "三、遇到的核心问题及解决方案", level=2)
add_text(doc, "核心问题          解决方案              技术要点", bold=True)
add_text(doc, "座位超卖          乐观锁机制            @Version注解实现版本控制")
add_text(doc, "跨域请求          Vite代理配置          开发环境proxy，生产环境Nginx")
add_text(doc, "路由404           Hash模式路由          根据环境动态选择路由模式")
add_text(doc, "构建超时          npmmirror镜像        npm install配置registry")

add_heading(doc, "四、系统思维的核心收获", level=2)
add_text(doc, "1. 整体性思维：", bold=True)
add_text(doc, "不再局限于单个模块，从整体架构角度思考问题。前后端分离需要统一的API规范和接口文档作为\"契约\"。")
add_text(doc, "")
add_text(doc, "2. 关联性思维：", bold=True)
add_text(doc, "理解组件之间的依赖关系。修改数据库结构时考虑对后端API的影响，后端改动时考虑对前端的影响。")
add_text(doc, "")
add_text(doc, "3. 反馈思维：", bold=True)
add_text(doc, "通过CI/CD实现快速反馈。每次提交触发完整构建，失败立即通知，形成\"小步快跑\"的迭代模式。")
add_text(doc, "")
add_text(doc, "4. 边界思维：", bold=True)
add_text(doc, "明确系统边界和外部依赖。AI服务不可用时返回友好提示，而不是直接报错崩溃。")

add_heading(doc, "五、未来改进方向", level=2)
add_text(doc, "5.1 高并发优化", bold=True)
add_bullet(doc, "引入Redis缓存热点数据（演出信息、座位状态）")
add_bullet(doc, "使用消息队列（RabbitMQ/Kafka）削峰，处理抢票高峰")
add_bullet(doc, "分布式Session管理，支持水平扩展")
add_text(doc, "")
add_text(doc, "5.2 安全性增强", bold=True)
add_bullet(doc, "JWT token刷新机制，支持长期登录")
add_bullet(doc, "API限流防刷，防止恶意抢票")
add_bullet(doc, "敏感数据加密存储（密码加盐哈希）")
add_text(doc, "")
add_text(doc, "5.3 智能化提升", bold=True)
add_bullet(doc, "智能座位推荐算法：根据用户偏好推荐最佳座位")
add_bullet(doc, "个性化演出推荐：基于用户历史购票行为推荐相似演出")
add_bullet(doc, "实时客服机器人：多轮对话，理解上下文语境")

add_heading(doc, "六、致谢", level=2)
add_text(doc, "感谢老师的悉心指导，在项目开发过程中给出了很多宝贵的建议。感谢同学们的帮助，在遇到问题时能够相互讨论、共同解决。")
add_text(doc, "")
add_text(doc, "通过这次项目，我不仅掌握了前后端开发、CI/CD部署等技术知识，更重要的是学会了用系统思维去分析和解决复杂问题。这将对我未来的学习和工作产生深远影响。")

# 结束语
doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run("===== 作业完成 =====")
run.bold = True
run.font.size = Pt(14)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 保存到新文件
doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_已完善.docx')
print("Document enhanced successfully!")
print("File saved as: 个人作业汇总报告_已完善.docx")
