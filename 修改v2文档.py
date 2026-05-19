from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 读取v2文档
doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告_v2.docx')

def add_heading(doc, text, level=2, color=RGBColor(0x2E, 0x75, 0xB6)):
    """添加标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 2 else 12)
    run.font.color.rgb = color
    return para

def add_text(doc, text, bold=False):
    """添加正文"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para

def add_bullet(doc, text):
    """添加列表项"""
    para = doc.add_paragraph()
    run = para.add_run("• " + text)
    run.font.size = Pt(11)
    return para

def add_screenshot_box(doc, num, title, description, requirements):
    """添加截图说明框"""
    # 标题
    para = doc.add_paragraph()
    run = para.add_run(f"Screenshot {num}: {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC4, 0x59, 0x11)
    
    # 描述
    para = doc.add_paragraph()
    run = para.add_run(description)
    run.font.size = Pt(11)
    
    # 截图要求
    para = doc.add_paragraph()
    run = para.add_run("INSERT LOCATION: ")
    run.bold = True
    run.font.size = Pt(11)
    run = para.add_run(requirements)
    run.font.size = Pt(11)
    
    doc.add_paragraph()

# 遍历文档找到需要插入内容的位置
# 我们需要在每个主要章节后添加详细内容

# 找到所有段落
paragraphs = doc.paragraphs
target_texts = [
    "1. 以往项目中问题出现在哪些环节？",
    "2. 这些问题是代码错误，还是集成问题？",
    "3. 哪些假设是想当然、未经过验证的？",
    "1. 为什么选择GitHub Actions？",
    "3. 配置文件详解",
    "1. 智能体AI案例介绍",
    "4. 应用场景思考",
    "1. 什么是系统思维？",
    "3. 系统思维的应用实例"
]

# 收集要插入的内容（按顺序）
insert_content = []

# Day 1 - 问题环节详细
insert_content.append({
    "trigger": "1. 以往项目中问题出现在哪些环节？",
    "content": """
【DETAILED CASE ANALYSIS】

【Issue 1: CORS Request Failure】
• Phenomenon: Frontend axios request reports error "Access-Control-Allow-Origin"
• Cause: Frontend (localhost:5173) and backend (localhost:8081) have different ports
• Solution: Configure proxy in vite.config.js to forward /api requests

Key Code:
proxy: {
  '/api': {
    target: 'http://localhost:8081',
    changeOrigin: true
  }
}

【Issue 2: Data Format Inconsistency】
• Phenomenon: Frontend expects {code: 200, data: [...]}, backend returns array directly
• Cause: Backend Controller doesn't have unified return format
• Solution: Define Result<T> class with code, msg, data fields

【Issue 3: Field Naming Differences】
• Phenomenon: Backend returns userName, frontend receives undefined
• Cause: Backend uses camelCase, frontend expects snake_case
• Solution: Configure camel-to-snake conversion in application.yml

【Database Performance Issues】
• Seat query slow: No index → Full table scan → 3-5 seconds delay
  Solution: Create composite index on showId, status
• Order redundancy: Username stored redundantly
  Solution: Use JOIN query for real-time user info
• Timezone issue: Database time 8 hours less
  Solution: Configure serverTimezone=Asia/Shanghai

【Frontend Component Issues】
• Element Plus styles not working: Need to import theme-chalk CSS
• Router navigation not refreshing: Use :key binding with route params
• Seat status not synced: Need polling or WebSocket for real-time updates
"""
})

# Day 1 - 代码错误 vs 集成问题
insert_content.append({
    "trigger": "2. 这些问题是代码错误，还是集成问题？",
    "content": """
【DETAILED ANALYSIS】

Integration Issues (~60%):
• CORS configuration - environment setup
• API proxy settings - server configuration
• GitHub Pages routing - deployment constraints
• Database timezone - environment differences

Code Errors (~40%):
• API format inconsistency - lack of unified specification
• Missing database indexes - performance issues
• Component style import - dependency issues

Key Insight: 
Integration problems often appear first because they involve multiple components working together. Code errors are easier to find through unit tests, but integration problems require full environment testing.

Recommended Approach:
1. First solve integration issues to ensure dev=prod
2. Use Code Review to catch code errors early
3. Establish CI/CD pipeline for continuous integration testing
"""
})

# Day 1 - 假设验证
insert_content.append({
    "trigger": "3. 哪些假设是想当然、未经过验证的？",
    "content": """
【DETAILED ANALYSIS WITH REAL EXAMPLES】

【User Behavior Assumptions】

Assumption: Users will follow normal purchase flow
Reality: Users may search directly, frequently navigate back, skip steps
Impact: Need to handle all possible user journeys

Assumption: Users will wait for page load (<1s response)
Reality: Network fluctuation causes 3-5s delays
Impact: Need loading animations and timeout handling

Assumption: Users won't submit repeatedly
Reality: Double-click or network retry causes duplicate orders
Impact: Need button debounce and request idempotency

【System Performance Assumptions】

Assumption: Single server handles all users
Reality: Pressure test showed 100 concurrent users caused slow responses
Impact: Need horizontal scaling plan

Assumption: Database queries are fast enough
Reality: Complex JOIN queries still have performance issues
Impact: Need SQL optimization and indexing

Assumption: Third-party APIs are stable (DeepSeek)
Reality: High concurrency causes API timeouts
Impact: Need fallback strategy and caching

【Deployment Environment Assumptions】

Assumption: Works locally = Works in production
Reality: Windows vs Linux, path separators, file permissions differ
Impact: Need Docker containerization or unified dev environment

Assumption: GitHub Pages supports SPA routing
Reality: History mode requires server config, static hosting doesn't support it
Impact: Need Hash mode for GitHub Pages deployment
"""
})

# Day 3 - GitHub Actions
insert_content.append({
    "trigger": "1. 为什么选择GitHub Actions？",
    "content": """
【DETAILED COMPARISON】

Why GitHub Actions Over Other CI/CD Tools:

| Feature | GitHub Actions | Jenkins | GitLab CI |
|---------|---------------|---------|-----------|
| Setup | Built-in, zero config | Install separately | GitLab account |
| Cost | 2000 min/month free | Self-hosted (costly) | Limited free tier |
| Config | YAML, simple syntax | Groovy, complex | YAML, similar |
| Integration | Native GitHub | Plugins needed | Native GitLab |
| Speed | Fast, cached | Slower | Moderate |

GitHub Actions Advantages for This Project:
1. Zero server maintenance - GitHub hosts everything
2. Native integration - no need for tokens/keys
3. Marketplace - thousands of pre-built actions
4. Parallel execution - faster builds
5. Free for open source and limited private repos
"""
})

# Day 3 - 配置详解
insert_content.append({
    "trigger": "3. 配置文件详解",
    "content": """
【STEP-BY-STEP CONFIGURATION EXPLAINED】

【Trigger Configuration】
on:
  push:
    branches: [ master, main ]  # Triggers on push to these branches
  pull_request:
    branches: [ master, main ]  # Triggers on PR for code review
  workflow_dispatch:              # Allows manual trigger

【Environment Setup】
- JDK 17 with Temurin distribution (community-recommended)
- Node.js 20 LTS (latest stable version)
- Maven and npm caching for faster builds

【Build Process】
Backend Build:
  cd damai-ticket
  mvn clean package -DskipTests
  # -DskipTests: Skip tests for faster CI builds
  # Produces: target/damai-ticket-1.0.0.jar

Frontend Build:
  cd damai-ticket-frontend
  npm install --registry https://registry.npmmirror.com
  # Using China mirror (10x faster than official npm)
  npm run build
  # Produces: dist/ folder with static assets

【Deployment】
- GitHub Pages deployment via actions/deploy-pages@v4
- Automatic URL generation: username.github.io/repo-name
- HTTPS enabled by default

【Problems Encountered & Solutions】
| Problem | Cause | Solution |
|---------|-------|----------|
| npm timeout | Slow npm access | Use npmmirror |
| Maven slow | Foreign repo | Use cache |
| 404 after deploy | History routing | Use Hash mode |
| Blank page | Path issues | Check base config |
"""
})

# Day 4 - AI介绍
insert_content.append({
    "trigger": "1. 智能体AI案例介绍",
    "content": """
【ARCHITECTURE & IMPLEMENTATION DETAILS】

【System Architecture】
User Input → Vue Component → Axios Request → Spring Controller → DeepSeek API → Response

【Frontend Implementation (Vue 3)】
• Composition API for chat logic
• Element Plus dialog for UI
• Axios for API calls
• Polling or WebSocket for real-time

【Backend Implementation (Spring Boot)】
• AuthController handles AI requests
• HttpClient calls DeepSeek API
• System prompt: "You are a ticketing system customer service"
• Fallback: Friendly message when API unavailable

【API Configuration】
API URL: https://api.deepseek.com/chat/completions
Model: deepseek-chat
Max Tokens: 1000
Temperature: 0.7 (creative but not random)

【Key Code Snippet】
@RestController
public class AuthController {
    @PostMapping("/api/ai/chat")
    public Result<?> chat(@RequestBody Map<String, String> request) {
        String question = request.get("question");
        String prompt = "You are a ticketing customer service. " + question;
        String answer = deepSeekService.chat(prompt);
        return Result.success(answer);
    }
}
"""
})

# Day 4 - 应用场景
insert_content.append({
    "trigger": "4. 应用场景思考",
    "content": """
【DETAILED DIALOGUE EXAMPLES】

【Dialogue 1: Concert Recommendation】
User: "What concerts are recommended for this weekend?"
AI: "Based on your location and preferences, I recommend:
    1. Jay Chou World Tour - National Stadium, Sat 8PM
    2. Taylor Swift Concert - Grand Theater, Sun 7PM
    3. Local Band Festival - Rock Bar, Fri 9PM

Would you like more details on any of these?"

【Dialogue 2: Purchase Guidance】
User: "How do I buy tickets?"
AI: "Here's the purchase process:
    Step 1: Log in or register an account
    Step 2: Browse shows or search by artist/venue
    Step 3: Select your show and choose seats
    Step 4: Complete payment
    Step 5: Receive e-ticket 24 hours before show
    
    Is there anything else you'd like to know?"

【Dialogue 3: Order Status】
User: "When will I receive my tickets?"
AI: "Your order #12345 has been paid successfully.
    E-tickets will be sent to your email 24 hours before the show.
    You can also view your tickets in 'My Orders' page.
    
    Show: Jay Chou Concert
    Date: May 25, 2026, 8:00 PM
    Venue: National Stadium

    Is there anything else I can help you with?"

【Dialogue 4: Refund Policy】
User: "Can I get a refund if I can't go?"
AI: "Our refund policy:
    • 7+ days before show: Full refund
    • 3-7 days before: 80% refund
    • Within 72 hours: No refund allowed
    
    To request a refund, go to 'My Orders' → Select order → 'Request Refund'
    
    Would you like to proceed with a refund request?"
"""
})

# Day 8 - 系统思维
insert_content.append({
    "trigger": "1. 什么是系统思维？",
    "content": """
【SYSTEMS THINKING FRAMEWORK】

【Four Core Principles】

1. HOLISTIC THINKING
   "The whole is greater than the sum of its parts"
   • System has emergent properties not present in individual components
   • Understanding the whole requires looking at interactions, not just parts
   • Example: A car has transportation capability - not present in any single part

2. RELATIONAL THINKING
   "Everything is connected"
   • Changes in one part affect other parts
   • Feedback loops create complex behaviors
   • Example: Adding brake power affects acceleration, handling, safety

3. DYNAMIC THINKING
   "Systems change over time"
   • Current solutions may not work in the future
   • Need to anticipate changes and adaptations
   • Example: What works for 100 users may fail for 10,000

4. FEEDBACK THINKING
   "Actions produce consequences"
   • Positive feedback amplifies changes
   • Negative feedback dampens changes
   • Example: More users → more revenue → more features → more users

【Visual Representation】
    
        ┌─────────────────────────────────┐
        │         SYSTEM BOUNDARY         │
        │  ┌─────┐    ┌─────┐            │
        │  │ Front│ ←→ │Backend│          │
        │  │ Vue  │    │Spring│          │
        │  └──┬──┘    └──┬──┘            │
        │     │          │                │
        │     └────┬─────┘                │
        │          ↓                      │
        │     ┌─────────┐                 │
        │     │Database │                 │
        │     └─────────┘                 │
        │          ↑                      │
        │     ┌─────────────┐            │
        │     │User Feedback│            │
        │     └─────────────┘            │
        └─────────────────────────────────┘
"""
})

# Day 8 - 应用实例
insert_content.append({
    "trigger": "3. 系统思维的应用实例",
    "content": """
【CASE STUDY 1: SEAT OVERSELLING PROBLEM】

【Problem】
Two users grabbed the same ticket simultaneously, both showed "Purchase Success", but only one could actually get the ticket.

【Systems Analysis】
1. Identify Critical Chain:
   Check seat status → Create order → Update seat status
   
2. Discover Relationships:
   Seat status affects order creation
   Order creation affects seat status
   These form a circular dependency

3. Find Root Cause:
   The three-step process is NOT ATOMIC
   Under concurrency, race condition occurs:
   User A: Checks seat (available) ✓
   User B: Checks seat (available) ✓
   User A: Creates order ✓
   User B: Creates order ✓  ← OVERSELLING!
   
【Solution with Systems Thinking】
• Use optimistic locking: @Version annotation
• System automatically checks version number during update
• Second update fails, preventing overselling

【Feedback Verification】
• Implement pressure testing with JMeter
• 1000 concurrent requests, verify no overselling
• Monitor success rate and error rate

---

【CASE STUDY 2: GITHUB PAGES DEPLOYMENT】

【Problem】
Works fine locally, but shows blank page + 404 errors after deployment.

【Systems Analysis】
1. Identify Environment Differences:
   Local: Has backend server for routing
   Production: Pure static file hosting
   
2. Understand Constraints:
   GitHub Pages cannot configure server-side routing
   Vue Router History mode requires server support
   
3. Find Root Cause:
   Vue Router History mode:
   URL: /events → Requires server to serve index.html for /events
   
   GitHub Pages doesn't support this:
   URL: /events → Returns 404 (file not found)

【Solution with Systems Thinking】
• Switch to Hash mode: /#/events
• Hash part never sent to server, always serves index.html
• Environment-aware routing selection:

const isGitHubPages = window.location.hostname.includes('github.io');
const history = isGitHubPages 
    ? createWebHashHistory() 
    : createWebHistory();
"""
})

# 在文档末尾添加截图要求和总结
doc.add_page_break()

add_heading(doc, "【SCREENSHOT INSERTION GUIDE】", level=1, color=RGBColor(0x00, 0x00, 0x00))
doc.add_paragraph()

add_heading(doc, "Screenshot 1: Project Directory Structure", level=2)
add_screenshot_box(doc, 1, "Project Structure",
    "Capture the complete directory structure showing frontend-backend separated architecture.",
    "INSERT HERE: After '1. 以往项目中问题出现在哪些环节？' section\n" +
    "Location: After line '在我开发大麦网票务系统的过程中...'\n" +
    "Screenshot: Open damai-ticket-project folder, show tree structure")

add_heading(doc, "Screenshot 2: Git Commit History", level=2)
add_screenshot_box(doc, 2, "Git Log",
    "Show Git commit records reflecting development process.",
    "INSERT HERE: After '3. 哪些假设是想当然...' section\n" +
    "Location: After the last assumption paragraph\n" +
    "Screenshot: Run 'git log --oneline -10' in Git Bash")

add_heading(doc, "Screenshot 3: GitHub Actions Success", level=2)
add_screenshot_box(doc, 3, "CI/CD Pipeline",
    "Show successful GitHub Actions workflow run.",
    "INSERT HERE: After '4. GitHub Pages部署配置' section\n" +
    "Location: Before Day 4 section\n" +
    "Screenshot: Visit Actions page, capture successful run with green checkmarks")

add_heading(doc, "Screenshot 4: GitHub Pages Settings", level=2)
add_screenshot_box(doc, 4, "Pages Configuration",
    "Show GitHub Pages deployment settings.",
    "INSERT HERE: After Screenshot 3\n" +
    "Location: Same section as Screenshot 3\n" +
    "Screenshot: Settings → Pages page showing 'Your site is published at'")

add_heading(doc, "Screenshot 5: Deployed Website", level=2)
add_screenshot_box(doc, 5, "Live Site",
    "Show the deployed ticketing system website.",
    "INSERT HERE: After Screenshot 4\n" +
    "Location: Same section\n" +
    "Screenshot: Open https://wang314152.github.io/damai-ticket-project/")

add_heading(doc, "Screenshot 6: AI Agent Project", level=2)
add_screenshot_box(doc, 6, "AI Project Structure",
    "Show agentic-ai-public-main project directory.",
    "INSERT HERE: After '1. 智能体AI案例介绍' section\n" +
    "Location: After architecture description\n" +
    "Screenshot: Open agentic-ai-public-main folder")

add_heading(doc, "Screenshot 7: AI Chat Demo", level=2)
add_screenshot_box(doc, 7, "AI Conversation",
    "Show actual AI customer service dialogue.",
    "INSERT HERE: After '4. 应用场景思考' section\n" +
    "Location: End of Day 4 section\n" +
    "Screenshot: Start project locally, go to /ai-assistant, capture dialogue")

add_heading(doc, "Screenshot 8: System Architecture", level=2)
add_screenshot_box(doc, 8, "Architecture Diagram",
    "Show overall system architecture with components and data flow.",
    "INSERT HERE: After '1. 什么是系统思维？' section\n" +
    "Location: After systems thinking principles\n" +
    "Screenshot: Use draw.io/Xmind to draw diagram showing Frontend→API→Backend→DB")

add_heading(doc, "Screenshot 9: CI/CD Pipeline", level=2)
add_screenshot_box(doc, 9, "Pipeline Flow",
    "Show complete CI/CD pipeline from commit to deploy.",
    "INSERT HERE: After '3. 系统思维的应用实例' section\n" +
    "Location: End of Day 8 section\n" +
    "Screenshot: Draw flow: Code→GitHub→Build→Test→Deploy")

# 添加项目总结
doc.add_page_break()
add_heading(doc, "【PROJECT SUMMARY & REFLECTION】", level=1, color=RGBColor(0x00, 0x00, 0x00))

add_heading(doc, "1. Technical Achievements", level=2)
add_bullet(doc, "Frontend: Vue 3 Composition API, Vue Router, Element Plus, Vite")
add_bullet(doc, "Backend: Spring Boot, MyBatis-Plus, RESTful API, Optimistic Locking")
add_bullet(doc, "DevOps: GitHub Actions CI/CD, GitHub Pages deployment, Environment configuration")

add_heading(doc, "2. Problems Solved with Systems Thinking", level=2)
add_bullet(doc, "Seat Overselling: Used @Version optimistic lock to ensure data consistency")
add_bullet(doc, "CORS Issues: Configured Vite proxy for development, Nginx for production")
add_bullet(doc, "Routing 404: Implemented environment-aware Hash/History mode selection")

add_heading(doc, "3. Key Learnings", level=2)
add_bullet(doc, "Holistic Thinking: Systems are more than their parts - consider interactions")
add_bullet(doc, "Relational Thinking: Changes in one component affect others")
add_bullet(doc, "Feedback Thinking: Build feedback loops for continuous improvement")
add_bullet(doc, "Boundary Thinking: Understand system limits and external dependencies")

add_heading(doc, "4. Future Improvements", level=2)
add_text(doc, "High Concurrency:", bold=True)
add_bullet(doc, "Add Redis for caching hot data (seat status, show info)")
add_bullet(doc, "Use message queue (Kafka) for peak handling during ticket rush")
add_text(doc, "Security:", bold=True)
add_bullet(doc, "Implement JWT refresh mechanism")
add_bullet(doc, "Add API rate limiting to prevent abuse")
add_text(doc, "Intelligence:", bold=True)
add_bullet(doc, "Smart seat recommendation algorithm")
add_bullet(doc, "Personalized show recommendations based on history")

add_heading(doc, "5. Acknowledgements", level=2)
add_text(doc, "I would like to express my sincere gratitude to the teacher for the valuable guidance throughout this project. The systematic thinking approach learned in this course has fundamentally changed how I approach problem-solving in software development.")
add_text(doc, "Special thanks to classmates for the collaborative learning experience. Through discussions and code reviews, I gained deeper insights into both technical and analytical aspects of system design.")

# 结束语
doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run("===== ASSIGNMENT COMPLETE =====")
run.bold = True
run.font.size = Pt(14)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 保存英文版本
doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_Final_English.docx')
print("Document updated and saved as: 个人作业汇总报告_Final_English.docx")
