from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 读取原文档
doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告.docx')

def add_heading(doc, text, level=2):
    """添加标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 2 else 12)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return para

def add_text(doc, text, bold=False):
    """添加正文"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para

def add_bullet(doc, text):
    """添加列表项"""
    para = doc.add_paragraph()
    run = para.add_run("• " + text)
    run.font.size = Pt(11)
    return para

# ==================== 在原文档末尾添加详细补充内容 ====================

# 添加分页符和标题
doc.add_paragraph()
add_heading(doc, "【DETAILED SUPPLEMENTS】", level=1)

# ==================== Day 1 Detailed Supplements ====================
doc.add_paragraph()
add_heading(doc, "Day 1 - Detailed Problem Analysis", level=1)

add_heading(doc, "1.1 Frontend-Backend Integration - Detailed Case Studies", level=2)

add_text(doc, "【Issue 1: CORS Request Failure】", bold=True)
add_bullet(doc, "Phenomenon: Frontend axios request reports error \"Access-Control-Allow-Origin\"")
add_bullet(doc, "Cause: Frontend dev server (localhost:5173) and backend server (localhost:8081) have different ports, triggering browser security restrictions")
add_bullet(doc, "Solution: Configure proxy in vite.config.js to forward /api requests to backend")

add_text(doc, "")
add_text(doc, "【Issue 2: Data Format Inconsistency】", bold=True)
add_bullet(doc, "Phenomenon: Frontend expects {code: 200, data: [...]}, but backend returns array directly")
add_bullet(doc, "Cause: Backend Controller doesn't have unified return format, no API specification agreed between frontend and backend")
add_bullet(doc, "Solution: Define unified response class Result<T> with code, msg, data fields")

add_text(doc, "")
add_text(doc, "【Issue 3: Field Naming Differences】", bold=True)
add_bullet(doc, "Phenomenon: Backend returns userName, frontend receives undefined")
add_bullet(doc, "Cause: Backend uses camelCase, frontend expects snake_case, JSON serialization config inconsistent")
add_bullet(doc, "Solution: Configure camel-to-snake conversion in application.yml or handle uniformly in frontend")

add_heading(doc, "1.2 Database Design - Performance Issues", level=2)
add_bullet(doc, "Seat table query slow: No index created, 10,000 records scanned causing 3-5 second delay")
add_bullet(doc, "  Solution: Create composite index on showId and status fields")
add_bullet(doc, "Order table redundancy: Username stored redundantly, historical orders cannot display latest info")
add_bullet(doc, "  Solution: Use JOIN query to get latest user info in real-time")
add_bullet(doc, "Timezone issue: Database time is 8 hours less than actual time")
add_bullet(doc, "  Solution: Configure serverTimezone=Asia/Shanghai in application.yml")

add_heading(doc, "3.1 User Behavior Assumptions - Detailed Analysis", level=2)
add_text(doc, "Assumption 1: Users will follow normal purchase flow (browse → select seat → pay)", bold=True)
add_text(doc, "Actual situation: Users may search directly, frequently go back and forth, need to consider diverse operation paths", False)
add_text(doc, "")
add_text(doc, "Assumption 2: Users will wait for page to load (interface responds <1 second)", bold=True)
add_text(doc, "Actual situation: May wait 3-5 seconds during network fluctuation, need loading animation and timeout handling", False)
add_text(doc, "")
add_text(doc, "Assumption 3: Users won't submit repeatedly", bold=True)
add_text(doc, "Actual situation: Users may double-click or network retry causing duplicate orders, need duplicate prevention mechanism", False)
add_text(doc, "Solution: Button debounce + Request idempotency design", False)

# ==================== Day 3 Detailed Supplements ====================
doc.add_paragraph()
add_heading(doc, "Day 3 - CI/CD Configuration Details", level=1)

add_heading(doc, "4.2 Build Steps - Step-by-Step Explanation", level=2)
add_text(doc, "1. Checkout code: Use actions/checkout@v4 to pull code", bold=True)
add_text(doc, "   This is the first step of GitHub Actions, essential", False)
add_text(doc, "")
add_text(doc, "2. Cache Maven packages: Cache Maven dependencies to speed up builds", bold=True)
add_text(doc, "   Key includes pom.xml hash, cache updates automatically when pom changes", False)
add_text(doc, "   Can save 2-3 minutes of dependency download time", False)
add_text(doc, "")
add_text(doc, "3. Set up JDK 17: Install Java 17 runtime environment", bold=True)
add_text(doc, "   Use temurin distribution (community-recommended OpenJDK distribution)", False)
add_text(doc, "   Faster and more stable than default adoptium", False)
add_text(doc, "")
add_text(doc, "4. Build Backend: cd damai-ticket && mvn clean package -DskipTests", bold=True)
add_text(doc, "   -DskipTests skips tests to speed up CI build", False)
add_text(doc, "")
add_text(doc, "5. Cache npm packages: Cache npm dependencies to speed up builds", bold=True)
add_text(doc, "   Based on package-lock.json hash", False)
add_text(doc, "")
add_text(doc, "6. Set up Node.js: Install Node.js 20 LTS version", bold=True)
add_text(doc, "   LTS version is more stable, 20 is the latest long-term support version", False)
add_text(doc, "")
add_text(doc, "7. Build Frontend: npm install (using npmmirror) + npm run build", bold=True)
add_text(doc, "   npmmirror is a China mirror, more than 10x faster than official npm", False)
add_text(doc, "")
add_text(doc, "8. Deploy to GitHub Pages: Upload dist directory, use actions/deploy-pages@v4", bold=True)

# ==================== Day 4 Detailed Supplements ====================
doc.add_paragraph()
add_heading(doc, "Day 4 - AI Application Scenarios", level=1)

add_heading(doc, "4.1 AI Application in Ticketing System - Specific Dialogue Examples", level=2)
add_text(doc, "【Dialogue 1: Intelligent Customer Service Q&A】", bold=True)
add_text(doc, "User: \"What concerts are recommended for this weekend?\"", False)
add_text(doc, "AI: \"Based on your location and preferences, I recommend the following concerts...\"", False)
add_text(doc, "")
add_text(doc, "【Dialogue 2: Purchase Guide】", bold=True)
add_text(doc, "User: \"How to buy tickets?\"", False)
add_text(doc, "AI: \"Purchase process: 1. Log in 2. Select show 3. Select seat 4. Pay\"", False)
add_text(doc, "")
add_text(doc, "【Dialogue 3: Order Inquiry】", bold=True)
add_text(doc, "User: \"When will my order be delivered?\"", False)
add_text(doc, "AI: \"Your order has been paid successfully. E-tickets will be sent 24 hours before the show\"", False)
add_text(doc, "")
add_text(doc, "【Dialogue 4: Refund Policy】", bold=True)
add_text(doc, "User: \"Can I get a refund?\"", False)
add_text(doc, "AI: \"Refunds are not supported within 72 hours before the show. See refund policy for details\"", False)

# ==================== Day 8 Detailed Supplements ====================
doc.add_paragraph()
add_heading(doc, "Day 8 - Systems Thinking Case Studies", level=1)

add_heading(doc, "3.1 Seat Overselling Problem - Systems Thinking Application", level=2)
add_text(doc, "【Problem Background】", bold=True)
add_text(doc, "User reported: Two people grabbed the same ticket simultaneously, both showed purchase success, but only one could actually get the ticket.", False)
add_text(doc, "")
add_text(doc, "【Systems Analysis】", bold=True)
add_bullet(doc, "Identify critical chain: Check seat status → Create order → Update seat status")
add_bullet(doc, "Discover relationships: Seat status and order creation affect each other, modification order is critical")
add_bullet(doc, "Find root cause: Check→Create→Update is not atomic, race condition occurs under concurrency")
add_text(doc, "")
add_text(doc, "【Solution】", bold=True)
add_bullet(doc, "Use MyBatis-Plus optimistic lock mechanism: @Version annotation handles version control automatically")
add_bullet(doc, "Automatically check version number during update, prevent data inconsistency from concurrent modifications")
add_bullet(doc, "Verify solution effectiveness through pressure testing")

add_text(doc, "")
add_heading(doc, "3.2 GitHub Pages Deployment Problem - Systems Thinking Application", level=2)
add_text(doc, "【Problem Background】", bold=True)
add_text(doc, "Works fine locally, but shows blank page after deploying to GitHub Pages. Clicking navigation links shows 404 error.", False)
add_text(doc, "")
add_text(doc, "【Systems Analysis】", bold=True)
add_bullet(doc, "Identify environment differences: Local has backend server to handle routing, production is pure static server")
add_bullet(doc, "Understand constraints: Static hosting cannot configure server-side routing rules")
add_bullet(doc, "Find root cause: Vue Router History mode depends on server configuration, GitHub Pages doesn't support it")
add_text(doc, "")
add_text(doc, "【Solution】", bold=True)
add_bullet(doc, "Switch to Hash mode routing: Route information stored after # in URL")
add_bullet(doc, "Code implementation: Select routing mode automatically based on environment")
add_bullet(doc, "const history = isGitHubPages ? createWebHashHistory() : createWebHistory();")

# ==================== Project Summary ====================
doc.add_paragraph()
add_heading(doc, "【PROJECT SUMMARY & OUTLOOK】", level=1)

add_heading(doc, "1. Project Overview", level=2)
add_text(doc, "This project is a frontend-backend separated online ticketing system, simulating Damai.com's ticket purchasing process. Using Vue 3 + Spring Boot tech stack, it implements core features like show browsing, seat selection, order management, and AI customer service, with a complete CI/CD automated deployment pipeline.")

add_heading(doc, "2. Technical Achievements", level=2)
add_text(doc, "2.1 Frontend Technologies", bold=True)
add_bullet(doc, "Vue 3 Composition API: Master component logic reuse, farewell to Options API complexity")
add_bullet(doc, "Vue Router: Deep understanding of SPA routing principles (History vs Hash mode)")
add_bullet(doc, "Element Plus: Component library usage and theme customization, improving development efficiency")
add_bullet(doc, "Vite: Experience modern build tools' fast development (hot reload in seconds)")
add_text(doc, "")
add_text(doc, "2.2 Backend Technologies", bold=True)
add_bullet(doc, "Spring Boot: Build enterprise application from scratch, understand auto-configuration principles")
add_bullet(doc, "MyBatis-Plus: Master ORM framework and optimistic lock mechanism, solve concurrency problems")
add_bullet(doc, "RESTful API: Interface design specifications, unified response format code/msg/data")
add_text(doc, "")
add_text(doc, "2.3 DevOps Technologies", bold=True)
add_bullet(doc, "GitHub Actions: CI/CD pipeline configuration, achieve automated build and deployment")
add_bullet(doc, "GitHub Pages: Static website hosting, free demo environment")
add_bullet(doc, "Environment configuration management: Dev/prod environment separation, Configuration as Code")

add_heading(doc, "3. Core Problems Encountered & Solutions", level=2)
add_text(doc, "Core Problem          Solution                  Technical Key Points", bold=True)
add_text(doc, "Seat overselling      Optimistic lock          @Version annotation for version control", False)
add_text(doc, "CORS issues          Vite proxy config        Dev proxy, prod Nginx", False)
add_text(doc, "Routing 404          Hash mode routing        Dynamic routing selection by environment", False)
add_text(doc, "Build timeout        npmmirror               npm registry configuration", False)

add_heading(doc, "4. Core Achievements in Systems Thinking", level=2)
add_text(doc, "1. Holistic Thinking:", bold=True)
add_text(doc, "No longer limited to single modules, think from overall architecture perspective. Frontend-backend separation requires unified API specifications and interface documents as \"contracts\".", False)
add_text(doc, "")
add_text(doc, "2. Relational Thinking:", bold=True)
add_text(doc, "Understand dependencies between components. Consider impact on backend API when modifying database structure; consider impact on frontend when modifying backend.", False)
add_text(doc, "")
add_text(doc, "3. Feedback Thinking:", bold=True)
add_text(doc, "Achieve rapid feedback through CI/CD. Each commit triggers complete build, failure notifies immediately, forming \"small steps, fast iterations\" model.", False)
add_text(doc, "")
add_text(doc, "4. Boundary Thinking:", bold=True)
add_text(doc, "Clarify system boundaries and external dependencies. Return friendly messages when AI service is unavailable, instead of crashing directly.", False)

add_heading(doc, "5. Future Improvement Directions", level=2)
add_text(doc, "5.1 High Concurrency Optimization", bold=True)
add_bullet(doc, "Introduce Redis to cache hot data (show info, seat status)")
add_bullet(doc, "Use message queue (RabbitMQ/Kafka) for peak shaving, handle ticket rush peaks")
add_bullet(doc, "Distributed session management, support horizontal scaling")
add_text(doc, "")
add_text(doc, "5.2 Security Enhancement", bold=True)
add_bullet(doc, "JWT token refresh mechanism, support long-term login")
add_bullet(doc, "API rate limiting and anti-brush, prevent malicious ticket grabbing")
add_bullet(doc, "Sensitive data encrypted storage (password salt hashing)")
add_text(doc, "")
add_text(doc, "5.3 Intelligence Enhancement", bold=True)
add_bullet(doc, "Intelligent seat recommendation algorithm: Recommend best seats based on user preferences")
add_bullet(doc, "Personalized show recommendation: Based on user's historical purchase behavior")
add_bullet(doc, "Real-time customer service robot: Multi-turn dialogue, understand context")

add_heading(doc, "6. Acknowledgements", level=2)
add_text(doc, "Thanks to the teacher for your careful guidance and valuable suggestions throughout the project development. Thanks to classmates for their help, we could discuss and solve problems together when encountered.", False)
add_text(doc, "")
add_text(doc, "Through this project, I not only mastered frontend-backend development and CI/CD deployment technical knowledge, but more importantly, learned to use systems thinking to analyze and solve complex problems. This will have a profound impact on my future learning and work.", False)

# ==================== Screenshot Requirements ====================
doc.add_paragraph()
add_heading(doc, "【SCREENSHOT REQUIREMENTS】", level=1)

add_heading(doc, "Screenshot 1: Project Directory Structure", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Capture the complete directory structure of the project, showing the frontend-backend separated architecture.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Open the project root folder (damai-ticket-project)", False)
add_text(doc, "• Show all main directories: damai-ticket/, damai-ticket-frontend/, .github/, agentic-ai-public-main/", False)
add_text(doc, "• Capture the tree view if possible", False)

add_heading(doc, "Screenshot 2: Git Commit History", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the Git commit records of the project, reflecting the problem-solving process during development.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Open Git Bash in the project folder", False)
add_text(doc, "• Run command: git log --oneline -10", False)
add_text(doc, "• Capture the commit history showing recent 10 commits", False)

add_heading(doc, "Screenshot 3: GitHub Actions Run Success", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show a successful GitHub Actions CI/CD pipeline run.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Visit: https://github.com/wang314152/damai-ticket-project/actions", False)
add_text(doc, "• Capture the most recent successful workflow run", False)
add_text(doc, "• Show all steps (checkout, build, deploy) with green checkmarks", False)
add_text(doc, "• Include the execution time for each step", False)

add_heading(doc, "Screenshot 4: GitHub Pages Settings", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the GitHub Pages deployment configuration page.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Go to Settings → Pages in your repository", False)
add_text(doc, "• Show: Source (Deploy from a branch)", False)
add_text(doc, "• Show: Branch (master / (root))", False)
add_text(doc, "• Show the published URL if available", False)

add_heading(doc, "Screenshot 5: GitHub Pages Website", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the deployed website effect on GitHub Pages.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Open: https://wang314152.github.io/damai-ticket-project/", False)
add_text(doc, "• Capture the homepage showing the ticketing system UI", False)
add_text(doc, "• Show the navigation menu and show list", False)

add_heading(doc, "Screenshot 6: AI Agent Project Structure", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the directory structure of the agentic-ai-public-main project.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Open the agentic-ai-public-main folder", False)
add_text(doc, "• Show directories: src/, static/, templates/, docker/", False)
add_text(doc, "• Show main files: main.py, requirements.txt, README.md", False)

add_heading(doc, "Screenshot 7: AI Chat Demo", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the actual AI customer service dialogue effect.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Start the project locally (backend on port 8081, frontend on port 5173)", False)
add_text(doc, "• Go to: http://localhost:5173/ai-assistant", False)
add_text(doc, "• Log in and ask some questions", False)
add_text(doc, "• Capture a dialogue showing AI responses", False)

add_heading(doc, "Screenshot 8: System Architecture Diagram", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the overall technical architecture of the Damai Ticketing System.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Use draw.io, Xmind, or any diagram tool", False)
add_text(doc, "• Show: Frontend (Vue) → API → Backend (Spring Boot) → Database", False)
add_text(doc, "• Include: AI service, CI/CD pipeline, GitHub Pages", False)
add_text(doc, "• Use arrows to show data flow direction", False)

add_heading(doc, "Screenshot 9: CI/CD Pipeline Diagram", level=2)
add_text(doc, "Description:", bold=True)
add_text(doc, "Show the complete CI/CD pipeline flow from code commit to deployment.")
add_text(doc, "")
add_text(doc, "Requirements:", bold=True)
add_text(doc, "• Draw a flow diagram showing:", False)
add_text(doc, "  Code Commit → GitHub → Trigger → Build Backend → Build Frontend → Deploy", False)
add_text(doc, "• Or screenshot from GitHub Actions showing the pipeline steps", False)
add_text(doc, "• Include step names and execution times if possible", False)

# End
doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run("===== ASSIGNMENT COMPLETE =====")
run.bold = True
run.font.size = Pt(14)
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save English version
doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_English.docx')
print("English document generated successfully!")
print("File saved as: 个人作业汇总报告_English.docx")
