from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

# 读取原文档
doc = Document('d:/2026/集成/damai-ticket-project/个人作业汇总报告.docx')

# 获取所有段落
paragraphs = doc.paragraphs

# 找到需要修改的段落并更新内容
for i, para in enumerate(paragraphs):
    text = para.text.strip()
    
    # Day 1 - 完善问题环节描述
    if "1.1 前后端联调环节" in text and i > 0:
        # 这个段落后面添加详细内容
        pass
    
    # 更新总结段落
    if text == "总结" or "总结" in text:
        para.clear()
        run = para.add_run("总结")
        run.bold = True
        run.font.size = Pt(14)
    
    # 完善系统思维反思
    if "4. 系统思维的收获与反思" in text:
        # 找到后续段落添加内容
        pass

# 在文档末尾添加项目总结（通过添加新段落）
# 找到最后一个非空段落，在其后添加总结

# 添加项目总结部分
summary_paragraphs = [
    ("项目总结与展望", True, 16),
    ("", False, 11),
    ("一、项目概述", True, 14),
    ("本项目是一个前后端分离的在线票务系统，模拟大麦网的演出票务购买流程。采用Vue 3 + Spring Boot技术栈，实现了演出浏览、座位选择、订单管理、AI智能客服等核心功能，并搭建了完整的CI/CD自动化部署流程。", False, 11),
    ("", False, 11),
    ("二、技术收获", True, 14),
    ("", False, 11),
    ("2.1 前端技术", True, 12),
    ("• Vue 3 Composition API：掌握组件逻辑复用，告别Options API的繁琐", False, 11),
    ("• Vue Router：深入理解SPA路由原理（History vs Hash模式）", False, 11),
    ("• Element Plus：组件库使用和主题定制，提升开发效率", False, 11),
    ("• Vite：体验到现代构建工具的快速开发体验（热更新秒级响应）", False, 11),
    ("", False, 11),
    ("2.2 后端技术", True, 12),
    ("• Spring Boot：从零搭建企业级应用，理解自动配置原理", False, 11),
    ("• MyBatis-Plus：掌握ORM框架和乐观锁机制，解决并发问题", False, 11),
    ("• RESTful API：接口设计规范，统一响应格式code/msg/data", False, 11),
    ("", False, 11),
    ("2.3 DevOps技术", True, 12),
    ("• GitHub Actions：CI/CD流水线配置，实现自动化构建部署", False, 11),
    ("• GitHub Pages：静态网站托管，免费的演示环境", False, 11),
    ("• 环境配置管理：开发/生产环境分离，配置即代码", False, 11),
    ("", False, 11),
    ("三、遇到的核心问题及解决方案", True, 14),
    ("", False, 11),
    ("核心问题          解决方案              技术要点", True, 11),
    ("座位超卖          乐观锁机制            @Version注解实现版本控制", False, 11),
    ("跨域请求          Vite代理配置          开发环境proxy，生产环境Nginx", False, 11),
    ("路由404           Hash模式路由          根据环境动态选择路由模式", False, 11),
    ("构建超时          npmmirror镜像        npm install配置registry", False, 11),
    ("", False, 11),
    ("四、系统思维的核心收获", True, 14),
    ("", False, 11),
    ("1. 整体性思维：", True, 11),
    ("不再局限于单个模块，从整体架构角度思考问题。前后端分离需要统一的API规范和接口文档作为\"契约\"。", False, 11),
    ("", False, 11),
    ("2. 关联性思维：", True, 11),
    ("理解组件之间的依赖关系。修改数据库结构时考虑对后端API的影响，后端改动时考虑对前端的影响。", False, 11),
    ("", False, 11),
    ("3. 反馈思维：", True, 11),
    ("通过CI/CD实现快速反馈。每次提交触发完整构建，失败立即通知，形成\"小步快跑\"的迭代模式。", False, 11),
    ("", False, 11),
    ("4. 边界思维：", True, 11),
    ("明确系统边界和外部依赖。AI服务不可用时返回友好提示，而不是直接报错崩溃。", False, 11),
    ("", False, 11),
    ("五、未来改进方向", True, 14),
    ("", False, 11),
    ("5.1 高并发优化", True, 12),
    ("• 引入Redis缓存热点数据（演出信息、座位状态）", False, 11),
    ("• 使用消息队列（RabbitMQ/Kafka）削峰，处理抢票高峰", False, 11),
    ("• 分布式Session管理，支持水平扩展", False, 11),
    ("", False, 11),
    ("5.2 安全性增强", True, 12),
    ("• JWT token刷新机制，支持长期登录", False, 11),
    ("• API限流防刷，防止恶意抢票", False, 11),
    ("• 敏感数据加密存储（密码加盐哈希）", False, 11),
    ("", False, 11),
    ("5.3 智能化提升", True, 12),
    ("• 智能座位推荐算法：根据用户偏好推荐最佳座位", False, 11),
    ("• 个性化演出推荐：基于用户历史购票行为推荐相似演出", False, 11),
    ("• 实时客服机器人：多轮对话，理解上下文语境", False, 11),
    ("", False, 11),
    ("六、致谢", True, 14),
    ("", False, 11),
    ("感谢老师的悉心指导，在项目开发过程中给出了很多宝贵的建议。感谢同学们的帮助，在遇到问题时能够相互讨论、共同解决。", False, 11),
    ("", False, 11),
    ("通过这次项目，我不仅掌握了前后端开发、CI/CD部署等技术知识，更重要的是学会了用系统思维去分析和解决复杂问题。这将对我未来的学习和工作产生深远影响。", False, 11),
    ("", False, 11),
    ("===== 作业完成 =====", True, 14),
]

# 在文档末尾添加总结
for text, is_bold, font_size in summary_paragraphs:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = is_bold
    run.font.size = Pt(font_size)
    if "作业完成" in text:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 保存修改后的文档
doc.save('d:/2026/集成/damai-ticket-project/个人作业汇总报告_已完善.docx')
print("文档修改成功！已保存为：个人作业汇总报告_已完善.docx")
